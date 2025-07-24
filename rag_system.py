# Enhanced rag_system.py - RAG System for Bank Information with Hybrid Integration
import os
import logging
from typing import List, Dict, Any, Optional
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from docx import Document
import pickle
from pathlib import Path
import re

logger = logging.getLogger(__name__)

class EnhancedBankRAGSystem:
    """Enhanced RAG System with improved document processing and hybrid integration."""
    
    def __init__(self, document_path: str = "Best_Bank.docx", model_name: str = "all-MiniLM-L6-v2"):
        """Initialize enhanced RAG system with document and embedding model."""
        self.document_path = document_path
        self.model = SentenceTransformer(model_name)
        self.index = None
        self.documents = []
        self.embeddings = None
        
        # Enhanced cache paths
        self.cache_dir = Path("rag_cache")
        self.cache_dir.mkdir(exist_ok=True)
        self.index_path = self.cache_dir / "faiss_index.bin"
        self.docs_path = self.cache_dir / "documents.pkl"
        self.embeddings_path = self.cache_dir / "embeddings.pkl"
        self.metadata_path = self.cache_dir / "metadata.pkl"
        
        # Enhanced document metadata
        self.document_metadata = []
        
        self.initialize_enhanced_rag_system()
    
    def load_document_enhanced(self) -> List[Dict[str, Any]]:
        """Enhanced document loading with metadata and better chunk management."""
        try:
            if not os.path.exists(self.document_path):
                logger.error(f"Document not found: {self.document_path}")
                return []
            
            doc = Document(self.document_path)
            processed_chunks = []
            
            # Extract text from paragraphs with enhanced metadata
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text and len(text) > 15:  # Skip very short chunks
                    # Enhanced chunk processing
                    chunk_data = {
                        "content": text,
                        "type": "paragraph",
                        "source_index": para_idx,
                        "length": len(text),
                        "word_count": len(text.split())
                    }
                    processed_chunks.append(chunk_data)
            
            # Extract text from tables with enhanced structure
            for table_idx, table in enumerate(doc.tables):
                table_content = []
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_content.append(row_text)
                
                if table_content:
                    # Combine table content
                    full_table_text = "\n".join(table_content)
                    chunk_data = {
                        "content": full_table_text,
                        "type": "table",
                        "source_index": table_idx,
                        "length": len(full_table_text),
                        "word_count": len(full_table_text.split()),
                        "rows": len(table_content)
                    }
                    processed_chunks.append(chunk_data)
            
            # Enhanced chunk splitting for large content
            final_chunks = []
            for chunk in processed_chunks:
                if chunk["length"] > 600:  # Split large chunks
                    sub_chunks = self._split_large_chunk(chunk)
                    final_chunks.extend(sub_chunks)
                else:
                    final_chunks.append(chunk)
            
            # Extract just the content for embeddings
            self.document_metadata = final_chunks
            text_chunks = [chunk["content"] for chunk in final_chunks]
            
            logger.info(f"Enhanced loading: {len(text_chunks)} chunks from {self.document_path}")
            logger.info(f"Chunk types: {self._get_chunk_type_summary()}")
            
            return text_chunks
            
        except Exception as e:
            logger.error(f"Error in enhanced document loading: {e}")
            return []
    
    def _split_large_chunk(self, chunk: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split large chunks while preserving metadata."""
        content = chunk["content"]
        sentences = re.split(r'[.!?]+', content)
        
        sub_chunks = []
        current_chunk = ""
        chunk_index = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            if len(current_chunk) + len(sentence) < 500:
                current_chunk += sentence + ". "
            else:
                if current_chunk.strip():
                    sub_chunk = chunk.copy()
                    sub_chunk.update({
                        "content": current_chunk.strip(),
                        "sub_index": chunk_index,
                        "length": len(current_chunk),
                        "word_count": len(current_chunk.split())
                    })
                    sub_chunks.append(sub_chunk)
                    chunk_index += 1
                
                current_chunk = sentence + ". "
        
        # Add remaining content
        if current_chunk.strip():
            sub_chunk = chunk.copy()
            sub_chunk.update({
                "content": current_chunk.strip(),
                "sub_index": chunk_index,
                "length": len(current_chunk),
                "word_count": len(current_chunk.split())
            })
            sub_chunks.append(sub_chunk)
        
        return sub_chunks
    
    def _get_chunk_type_summary(self) -> Dict[str, int]:
        """Get summary of chunk types for logging."""
        type_counts = {}
        for chunk in self.document_metadata:
            chunk_type = chunk.get("type", "unknown")
            type_counts[chunk_type] = type_counts.get(chunk_type, 0) + 1
        return type_counts
    
    def create_embeddings_enhanced(self, documents: List[str]) -> np.ndarray:
        """Create embeddings with enhanced error handling and progress tracking."""
        try:
            logger.info(f"Creating enhanced embeddings for {len(documents)} documents...")
            
            # Batch processing for better performance
            batch_size = 32
            all_embeddings = []
            
            for i in range(0, len(documents), batch_size):
                batch = documents[i:i + batch_size]
                batch_embeddings = self.model.encode(
                    batch, 
                    show_progress_bar=True if i == 0 else False,
                    convert_to_numpy=True
                )
                all_embeddings.append(batch_embeddings)
                
                if i % (batch_size * 10) == 0:  # Log every 10 batches
                    logger.info(f"Processed {min(i + batch_size, len(documents))}/{len(documents)} documents")
            
            embeddings = np.vstack(all_embeddings)
            logger.info(f"Enhanced embeddings created with shape: {embeddings.shape}")
            return embeddings
            
        except Exception as e:
            logger.error(f"Error creating enhanced embeddings: {e}")
            return np.array([])
    
    def build_faiss_index_enhanced(self, embeddings: np.ndarray) -> faiss.Index:
        """Build enhanced FAISS index with better configuration."""
        try:
            dimension = embeddings.shape[1]
            
            # Choose index type based on dataset size
            if len(embeddings) < 1000:
                # Use flat index for small datasets
                index = faiss.IndexFlatIP(dimension)
            else:
                # Use more sophisticated index for larger datasets
                nlist = min(100, len(embeddings) // 10)  # Number of clusters
                index = faiss.IndexIVFFlat(faiss.IndexFlatIP(dimension), dimension, nlist)
                
                # Train the index
                logger.info("Training FAISS index...")
                index.train(embeddings.astype(np.float32))
            
            # Normalize embeddings for cosine similarity
            faiss.normalize_L2(embeddings)
            
            # Add embeddings to index
            index.add(embeddings.astype(np.float32))
            
            logger.info(f"Enhanced FAISS index built with {index.ntotal} vectors")
            return index
            
        except Exception as e:
            logger.error(f"Error building enhanced FAISS index: {e}")
            return None
    
    def save_cache_enhanced(self):
        """Save enhanced cache with metadata."""
        try:
            if self.index:
                faiss.write_index(self.index, str(self.index_path))
            
            with open(self.docs_path, 'wb') as f:
                pickle.dump(self.documents, f)
            
            with open(self.embeddings_path, 'wb') as f:
                pickle.dump(self.embeddings, f)
            
            with open(self.metadata_path, 'wb') as f:
                pickle.dump(self.document_metadata, f)
            
            logger.info("Enhanced RAG cache saved successfully")
        except Exception as e:
            logger.error(f"Error saving enhanced cache: {e}")
    
    def load_cache_enhanced(self) -> bool:
        """Load enhanced cache with metadata validation."""
        try:
            if (self.index_path.exists() and 
                self.docs_path.exists() and 
                self.embeddings_path.exists() and
                self.metadata_path.exists()):
                
                # Check if document was modified after cache
                doc_mtime = os.path.getmtime(self.document_path)
                cache_mtime = os.path.getmtime(self.index_path)
                
                if doc_mtime > cache_mtime:
                    logger.info("Document updated since cache, rebuilding...")
                    return False
                
                self.index = faiss.read_index(str(self.index_path))
                
                with open(self.docs_path, 'rb') as f:
                    self.documents = pickle.load(f)
                
                with open(self.embeddings_path, 'rb') as f:
                    self.embeddings = pickle.load(f)
                
                with open(self.metadata_path, 'rb') as f:
                    self.document_metadata = pickle.load(f)
                
                logger.info(f"Enhanced RAG cache loaded: {len(self.documents)} documents, {self.index.ntotal} vectors")
                logger.info(f"Chunk types: {self._get_chunk_type_summary()}")
                return True
                
        except Exception as e:
            logger.error(f"Error loading enhanced cache: {e}")
        
        return False
    
    def initialize_enhanced_rag_system(self):
        """Initialize the enhanced RAG system."""
        try:
            # Try to load from cache first
            if self.load_cache_enhanced():
                return
            
            # Build from scratch
            logger.info("Building enhanced RAG system from scratch...")
            
            # Load document with enhanced processing
            self.documents = self.load_document_enhanced()
            if not self.documents:
                logger.error("No documents loaded!")
                return
            
            # Create embeddings with enhanced processing
            self.embeddings = self.create_embeddings_enhanced(self.documents)
            if self.embeddings.size == 0:
                logger.error("No embeddings created!")
                return
            
            # Build enhanced FAISS index
            self.index = self.build_faiss_index_enhanced(self.embeddings)
            if not self.index:
                logger.error("Failed to build enhanced FAISS index!")
                return
            
            # Save enhanced cache
            self.save_cache_enhanced()
            
            logger.info("Enhanced RAG system initialized successfully!")
            
        except Exception as e:
            logger.error(f"Error initializing enhanced RAG system: {e}")
    
    def search_relevant_chunks_enhanced(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Enhanced search for relevant document chunks with metadata."""
        try:
            if not self.index or not self.documents:
                logger.error("Enhanced RAG system not properly initialized")
                return []
            
            # Create query embedding
            query_embedding = self.model.encode([query])
            faiss.normalize_L2(query_embedding)
            
            # Search in FAISS index
            scores, indices = self.index.search(query_embedding.astype(np.float32), top_k)
            
            results = []
            for i, (score, idx) in enumerate(zip(scores[0], indices[0])):
                if idx < len(self.documents) and idx < len(self.document_metadata):
                    metadata = self.document_metadata[idx] if self.document_metadata else {}
                    results.append({
                        "chunk": self.documents[idx],
                        "score": float(score),
                        "rank": i + 1,
                        "metadata": metadata,
                        "type": metadata.get("type", "unknown"),
                        "word_count": metadata.get("word_count", 0)
                    })
            
            logger.info(f"Enhanced search: found {len(results)} relevant chunks for query: {query[:50]}...")
            return results
            
        except Exception as e:
            logger.error(f"Error in enhanced search: {e}")
            return []
    
    def generate_rag_response_enhanced(self, query: str, max_chunks: int = 4) -> Dict[str, Any]:
        """Generate enhanced RAG response with better context management."""
        try:
            # Search for relevant chunks
            relevant_chunks = self.search_relevant_chunks_enhanced(query, top_k=max_chunks)
            
            if not relevant_chunks:
                return {
                    "success": False,
                    "response": "I apologize, but I couldn't find relevant information about that in our bank documentation.",
                    "context_used": [],
                    "search_quality": "no_results"
                }
            
            # Enhanced relevance filtering
            good_chunks = [chunk for chunk in relevant_chunks if chunk["score"] > 0.25]
            
            if not good_chunks:
                return {
                    "success": False,
                    "response": "I don't have specific information about that. I can only help with bank-related queries from our official documentation.",
                    "context_used": [],
                    "search_quality": "low_relevance"
                }
            
            # Prioritize different chunk types
            prioritized_chunks = self._prioritize_chunks(good_chunks)
            
            # Combine context from relevant chunks
            context_parts = []
            for chunk in prioritized_chunks:
                chunk_content = chunk["chunk"]
                chunk_type = chunk.get("type", "paragraph")
                
                if chunk_type == "table":
                    context_parts.append(f"Table Information:\n{chunk_content}")
                else:
                    context_parts.append(chunk_content)
            
            context = "\n\n".join(context_parts)
            
            # Enhanced response metadata
            search_quality = self._assess_search_quality(good_chunks)
            
            return {
                "success": True,
                "context": context,
                "relevant_chunks": good_chunks,
                "query": query,
                "search_quality": search_quality,
                "chunk_types_used": [chunk.get("type", "unknown") for chunk in good_chunks],
                "total_word_count": sum(chunk.get("word_count", 0) for chunk in good_chunks)
            }
            
        except Exception as e:
            logger.error(f"Error generating enhanced RAG response: {e}")
            return {
                "success": False,
                "response": "I encountered an error while searching our documentation. Please try rephrasing your question.",
                "context_used": [],
                "search_quality": "error"
            }
    
    def _prioritize_chunks(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Prioritize chunks based on type and relevance."""
        # Sort by score first, then prioritize tables
        def chunk_priority(chunk):
            base_score = chunk["score"]
            if chunk.get("type") == "table":
                return base_score + 0.1  # Slight boost for tables
            return base_score
        
        return sorted(chunks, key=chunk_priority, reverse=True)
    
    def _assess_search_quality(self, chunks: List[Dict[str, Any]]) -> str:
        """Assess the quality of search results."""
        if not chunks:
            return "no_results"
        
        avg_score = sum(chunk["score"] for chunk in chunks) / len(chunks)
        
        if avg_score > 0.7:
            return "excellent"
        elif avg_score > 0.5:
            return "good"
        elif avg_score > 0.3:
            return "fair"
        else:
            return "poor"
    
    def is_bank_related_query_enhanced(self, query: str) -> bool:
        """Enhanced bank-related query detection with better patterns."""
        bank_keywords = [
            # Core banking terms
            "bank", "banking", "account", "balance", "transaction", "transfer", "deposit", "withdrawal",
            "loan", "credit", "debit", "card", "savings", "checking", "interest", "fee", "charge",
            
            # Services and operations
            "service", "branch", "atm", "online", "mobile", "app", "support", "help", "customer",
            "hours", "location", "contact", "phone", "email", "address", "statement", "report",
            
            # Financial products
            "mortgage", "investment", "insurance", "pension", "retirement", "business", "personal",
            "student", "auto", "home", "equity", "certificate", "money market", "fixed deposit",
            
            # Banking processes
            "application", "approval", "eligibility", "documentation", "requirement", "process",
            "policy", "terms", "conditions", "limit", "minimum", "maximum", "currency", "exchange",
            
            # Best Bank specific
            "best bank", "bestbank", "our bank", "this bank", "bank policy", "bank service",
            
            # Digital banking
            "internet banking", "mobile banking", "online transfer", "digital payment", "e-banking",
            
            # Customer service
            "complaint", "feedback", "suggestion", "problem", "issue", "resolution", "support ticket"
        ]
        
        query_lower = query.lower()
        
        # Direct keyword matching
        if any(keyword in query_lower for keyword in bank_keywords):
            return True
        
        # Pattern-based detection
        banking_patterns = [
            r'\b(how to|how can i|can i|where to)\s+.*\b(open|close|transfer|deposit|withdraw)',
            r'\b(what is|what are)\s+.*\b(rate|fee|charge|limit|policy)',
            r'\b(when is|when are)\s+.*\b(open|close|available)',
            r'\bmoney\s+(transfer|send|receive|exchange)',
            r'\baccount\s+(opening|closing|balance|statement)'
        ]
        
        for pattern in banking_patterns:
            if re.search(pattern, query_lower):
                return True
        
        return False
    
    def get_system_stats(self) -> Dict[str, Any]:
        """Get enhanced system statistics."""
        return {
            "documents_loaded": len(self.documents),
            "embeddings_shape": self.embeddings.shape if self.embeddings is not None else None,
            "index_total": self.index.ntotal if self.index else 0,
            "chunk_types": self._get_chunk_type_summary(),
            "cache_exists": all([
                self.index_path.exists(),
                self.docs_path.exists(),
                self.embeddings_path.exists(),
                self.metadata_path.exists()
            ]),
            "document_path": self.document_path,
            "model_name": self.model.get_sentence_embedding_dimension() if self.model else None
        }

# Initialize enhanced global RAG system instance
try:
    bank_rag = EnhancedBankRAGSystem()
    logger.info("Enhanced Bank RAG System initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize enhanced RAG system: {e}")
    bank_rag = None